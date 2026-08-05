from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

OUT = "/workspace/scratch/e5c7d9a5fab7/Intra_Rater_Reliability_Supplement_REVISED.docx"

BLUE = "2E5D7B"
DARK = "20323F"
MUTED = "5F6B73"
LIGHT = "EEF3F6"
PALE = "F7F9FA"
RED = "8A2D2D"
WHITE = "FFFFFF"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.78)
sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(0.82)
sec.right_margin = Inches(0.82)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.12

for name, size, before, after in [("Heading 1", 15, 14, 6), ("Heading 2", 12, 10, 4), ("Heading 3", 10.5, 7, 3)]:
    st = styles[name]
    st.font.name = "Arial"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(BLUE if name != "Heading 3" else DARK)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

caption = styles["Caption"]
caption.font.name = "Arial"
caption._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
caption.font.size = Pt(9)
caption.font.italic = False
caption.font.color.rgb = RGBColor.from_string(DARK)
caption.paragraph_format.space_before = Pt(4)
caption.paragraph_format.space_after = Pt(4)

if "Key finding" not in styles:
    styles.add_style("Key finding", WD_STYLE_TYPE.PARAGRAPH)
key = styles["Key finding"]
key.font.name = "Arial"
key._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
key._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
key.font.size = Pt(10.5)
key.font.bold = True
key.font.color.rgb = RGBColor.from_string(DARK)
key.paragraph_format.space_before = Pt(7)
key.paragraph_format.space_after = Pt(7)
key.paragraph_format.left_indent = Inches(0.16)
key.paragraph_format.right_indent = Inches(0.16)

def shade_paragraph(p, fill=LIGHT):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "16")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BLUE)
    borders.append(left)
    pPr.append(borders)

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_table_borders(table, color="C7D0D6", size="4"):
    tblPr = table._tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)
        borders.append(el)

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def add_table(headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.width = Inches(widths[j])
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.bold = True; r.font.color.rgb = RGBColor.from_string(WHITE); r.font.size = Pt(font_size)
        r.font.name = "Arial"
    set_repeat_table_header(table.rows[0])
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].width = Inches(widths[j])
            set_cell_margins(cells[j])
            if i % 2 == 1: set_cell_shading(cells[j], PALE)
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val)); r.font.name = "Arial"; r.font.size = Pt(font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table

def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)

def new_numbering_id():
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    lvl = OxmlElement("w:lvl"); lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); lvl.append(start)
    numFmt = OxmlElement("w:numFmt"); numFmt.set(qn("w:val"), "decimal"); lvl.append(numFmt)
    lvlText = OxmlElement("w:lvlText"); lvlText.set(qn("w:val"), "%1."); lvl.append(lvlText)
    suff = OxmlElement("w:suff"); suff.set(qn("w:val"), "tab"); lvl.append(suff)
    pPr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num"); tab.set(qn("w:pos"), "460"); tabs.append(tab); pPr.append(tabs)
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "460"); ind.set(qn("w:hanging"), "260"); pPr.append(ind)
    lvl.append(pPr); abstract.append(lvl); numbering.append(abstract)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId"); abs_ref.set(qn("w:val"), str(abstract_id)); num.append(abs_ref)
    numbering.append(num)
    return num_id

def add_num(text, num_id):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId"); numId.set(qn("w:val"), str(num_id))
    numPr.append(ilvl); numPr.append(numId); pPr.append(numPr)
    p.paragraph_format.left_indent = Inches(0.32)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)

# Running furniture
hp = sec.header.paragraphs[0]
hp.text = "VALIDATION SUPPLEMENT  |  MSc Dissertation 2026"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs:
    r.font.name = "Arial"; r.font.size = Pt(8); r.font.color.rgb = RGBColor.from_string(MUTED)
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("Does Where You Are Shape What You Get?  |  Nor Lakrimdi")
run.font.name = "Arial"; run.font.size = Pt(8); run.font.color.rgb = RGBColor.from_string(MUTED)

# Title block
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
r = p.add_run("VALIDATION SUPPLEMENT")
r.font.name = "Arial"; r.font.size = Pt(10); r.bold = True; r.font.color.rgb = RGBColor.from_string(BLUE)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Intra-Rater Repeat Coding Audit")
r.font.name = "Arial"; r.font.size = Pt(23); r.bold = True; r.font.color.rgb = RGBColor.from_string(DARK)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
r = p.add_run("Instrument drift, agreement estimates, and implications for the human reference labels")
r.font.name = "Arial"; r.font.size = Pt(12); r.font.color.rgb = RGBColor.from_string(MUTED)

p = doc.add_paragraph(style="Key finding")
p.add_run("Bottom line. ").bold = True
p.add_run("The repeat-coding exercise did not hold the coding instrument constant. It is therefore evidence of coding drift, not confirmation that the human labels are acceptably reliable. A codebook-aligned repeat assessment is required before the first-round labels can be treated as a stable validation reference.")
shade_paragraph(p)

doc.add_heading("1. Purpose and scope", level=1)
doc.add_paragraph(
    "This supplement audits a repeat-coding exercise conducted by the same researcher on a subset of a 160-response hold-out dataset. Forty responses were selected for repeat coding; one uncodeable record was excluded, leaving 39 matched pairs. The aim is to describe agreement, identify procedural failures, and state what these results can and cannot support."
)
doc.add_paragraph(
    "The exercise concerns consistency of human coding. It does not validate the automated measures used for RQ1-RQ3, does not estimate inter-rater reliability, and does not establish a performance ceiling for an automated classifier."
)

doc.add_heading("2. Audit of the repeat-coding procedure", level=1)
doc.add_heading("2.1 Non-equivalent actionability scales", level=2)
doc.add_paragraph(
    "Overall actionability was scored on a 0-2 scale in Round 1 and a 0-4 scale in Round 2. Because the response categories and thresholds were not invariant, a direct ordinal kappa would not estimate repeat application of the same instrument and is not reported."
)
doc.add_paragraph(
    "A post hoc binary diagnostic was calculated using Round 1 >= 2 and Round 2 >= 3. It produced 84.6% observed agreement and kappa = 0.421 (n = 39). This result is descriptive only: it combines genuine coder inconsistency with the consequences of changing the scale, and it must not be presented as a valid ordinal intra-rater reliability estimate."
)

doc.add_heading("2.2 Binary-code definition drift", level=2)
doc.add_paragraph(
    "The marginal positive rates shifted markedly between rounds for several binary variables. These shifts are compatible with a more liberal interpretation in Round 2, but the numerical outputs alone cannot establish why each disagreement occurred. Explanations such as counting general supportive language as a coping step, or treating any telephone number as a crisis action, require a documented case-level audit against a frozen codebook."
)

doc.add_heading("2.3 Reference status", level=2)
doc.add_paragraph(
    "Neither coding round is intrinsically a gold standard. Accordingly, discordant pairs are described below as Round-2 positive shifts (R1 = 0, R2 = 1) and Round-2 negative shifts (R1 = 1, R2 = 0), not as false positives or false negatives. Round 1 may be retained provisionally for existing analyses because it was the pre-existing label set, but its correctness cannot be inferred merely from being stricter."
)

doc.add_heading("3. Results", level=1)
p = doc.add_paragraph("Table 1. Agreement and marginal distributions for binary codes (n = 39 matched pairs)", style="Caption")
rows = [
    ("coping_step", "0.176", "69.2%", "64.1%", "94.9%", "12", "0"),
    ("professional_help", "0.530", "92.3%", "89.7%", "92.3%", "2", "1"),
    ("social_support", "0.545", "84.6%", "71.8%", "87.2%", "6", "0"),
    ("crisis_action", "0.421", "69.2%", "25.6%", "56.4%", "12", "0"),
    ("follow_up", "0.000", "51.3%", "0.0%", "48.7%", "19", "0"),
    ("surface_localisation", "0.000", "89.7%", "89.7%", "100.0%", "4", "0"),
    ("verified_localisation", "0.552", "76.9%", "46.2%", "69.2%", "9", "0"),
]
add_table(["Measure", "Cohen's kappa", "Observed agreement", "R1 positive", "R2 positive", "0 to 1 shifts", "1 to 0 shifts"], rows, [1.55, .7, .85, .75, .75, .72, .72], 7.6)
doc.add_paragraph(
    "Note. Unweighted Cohen's kappa is reported for binary variables. Kappa depends on the marginal distributions and may be low when a category is nearly constant; observed agreement and both rounds' positive rates are therefore shown alongside it. Shift counts describe direction only and do not assign correctness."
).style = "Caption"

doc.add_heading("3.1 Interpretation of the observed patterns", level=2)
add_bullet("coping_step shows weak consistency and a large positive-rate shift (64.1% to 94.9%). This code cannot presently support a claim of stable single-coder application.")
add_bullet("professional_help and social_support have high observed agreement, but their kappa estimates remain moderate and specificity-like quantities are unstable because negatives are uncommon.")
add_bullet("crisis_action has only 69.2% agreement, with 12 Round-2 positive shifts. Given its relevance to crisis-contact analyses, these cases require codebook-based review before substantive interpretation.")
add_bullet("follow_up demonstrates severe instrument drift: Round 1 contains no positive cases, whereas Round 2 contains 19. Kappa is uninformative as a reliability summary when one round is constant; the code should not be used in reliability-based claims in its current form.")
add_bullet("surface_localisation has 89.7% observed agreement but kappa = 0 because Round 2 is uniformly positive. The high agreement is driven by prevalence and does not demonstrate discrimination between present and absent localisation.")
add_bullet("verified_localisation has the highest kappa (0.552) but also nine Round-2 positive shifts. This is not sufficient to certify the verification audit as reliable without case-level evidence and a stable verification protocol.")

doc.add_heading("4. Statistical reporting corrections", level=1)
reporting_nums = new_numbering_id()
add_num("Do not interpret kappa = 0 as proof that coding is statistically indistinguishable from chance. Kappa is a chance-corrected descriptive agreement coefficient, and here its value is strongly affected by constant or near-constant marginals.", reporting_nums)
add_num("Do not use sensitivity, specificity, or F1 as primary intra-rater measures unless one round has been independently justified as the reference. Those measures are asymmetric; swapping the designated reference can change their interpretation.", reporting_nums)
add_num("Do not infer a mechanism for disagreement from a confusion matrix. Claims about rule reinterpretation must be supported by a logged, case-level discrepancy review conducted against the codebook.", reporting_nums)
add_num("Do not use conventional verbal kappa bands as an acceptability test. If reported, they are descriptive conventions only; acceptability must be tied to a threshold specified before examining the results and justified for the measurement purpose.", reporting_nums)
add_num("Do not claim that intra-rater agreement sets the maximum achievable performance of an automated coder. Human repeatability and model-to-human agreement estimate different relationships and are not bounded in that way in a finite sample.", reporting_nums)

doc.add_heading("5. Consistency with the dissertation methodology", level=1)
p = doc.add_paragraph(style="Key finding")
p.add_run("Consistency verdict. ").bold = True
p.add_run("The original supplement is inconsistent with the dissertation's stated validation framework because it overstates human-label reliability and links the repeat-coding results to automated outcomes that were constructed and validated differently.")
shade_paragraph(p, "F3EDED")

add_bullet("The dissertation defines actionability_v2 as a 0-5 automated composite, whereas the repeat-coding exercise compared human actionability scales of 0-2 and 0-4. The repeat exercise cannot validate the 0-5 composite or its reported development-set F1 of 0.744.")
add_bullet("The dissertation already states that actionability_v2 and localisation_v2c were rebuilt and rescored on the same 112 labels, so their reported F1 values are development-set estimates, not independent hold-out confirmation. This supplement must preserve that caveat.")
add_bullet("The dissertation states that all human labels came from a single coder and that inter-rater reliability was not computed. This supplement may add an intra-rater diagnostic, but it cannot remove that limitation because the repeat instrument drifted.")
add_bullet("The repeat variables professional_help, social_support, coping_step, crisis_action, and follow_up are not identical to the five automated components of actionability_v2. Component-level repeat agreement therefore cannot be used to declare the H1 composite reliable.")
add_bullet("The crisis-contact fabrication audit requires verification of extracted contacts. verified_localisation repeat agreement is relevant to audit quality, but it does not by itself establish the accuracy of the fabricated-contact classifications or rescue the exploratory S1 finding.")

doc.add_heading("6. Implications for current findings", level=1)
doc.add_heading("6.1 Confirmatory findings", level=2)
doc.add_paragraph(
    "The repeat-coding results do not overturn the numerical models already fitted to the 1,120-response dataset, but they weaken the evidential basis for treating the human labels as a stable measurement reference. RQ1 and RQ2 should therefore retain the dissertation's existing qualification that the revised automated measures were evaluated on development labels and require independent hold-out confirmation. The repeat-coding exercise supplies no basis for upgrading those findings to confirmed results."
)
doc.add_heading("6.2 Exploratory crisis-contact finding", level=2)
doc.add_paragraph(
    "The visibly suspicious or unverifiable crisis-contact finding remains an exploratory audit result. Its credibility depends primarily on a complete, source-documented contact-verification table and reproducible case classifications, not on the binary repeat-coding kappas reported here. Any geographic comparison must distinguish verified, incorrect, suspicious, and unknown contacts and avoid treating absence of evidence as proof of fabrication."
)

doc.add_heading("7. Required corrective procedure", level=1)
doc.add_paragraph("Before the repeat assessment can support reliability claims, complete one clean recoding of the 39 valid cases:")
procedure_nums = new_numbering_id()
add_num("Freeze one codebook version word-for-word, including the original 0-2 actionability scale and operational examples for every binary field.", procedure_nums)
add_num("Present the 39 responses in a newly randomised order, with all previous labels and derived variables hidden.", procedure_nums)
add_num("Record coder identity, date, washout interval, workbook version, exclusions, and any rule queries raised before labels are revealed.", procedure_nums)
add_num("Do not reconcile labels during coding. Preserve the untouched recoding file and create a separate discrepancy log after the reliability statistics are frozen.", procedure_nums)
add_num("Report quadratic-weighted kappa for 0-2 actionability; unweighted kappa plus observed agreement and marginal prevalence for binary fields; and bootstrap confidence intervals where feasible.", procedure_nums)
add_num("Interpret acceptability against a prospectively stated criterion. If the aligned repeat still shows weak or systematically directional agreement, narrow or withdraw the affected measure rather than selecting the more favourable round.", procedure_nums)

doc.add_heading("8. Submission-ready reporting language", level=1)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.2)
p.paragraph_format.right_indent = Inches(0.2)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(8)
p.add_run(
    "A random subset of 40 responses from the 160-response hold-out dataset was recoded by the same researcher after a washout period, with first-round labels concealed; one uncodeable record was excluded, leaving 39 matched pairs. Audit of the repeat exercise identified a procedural failure: overall actionability was scored on a 0-4 scale in the repeat round rather than the original 0-2 scale, precluding a valid ordinal intra-rater estimate. The binary codes also showed substantial marginal shifts between rounds. Observed agreement ranged from 51.3% to 92.3%, while Cohen's kappa ranged from 0.000 to 0.552. Particularly large Round-2 positive shifts occurred for coping steps, crisis actions, follow-up, surface localisation, and verified localisation. Because the coding instrument was not invariant across rounds, these results are interpreted as evidence of coding drift rather than confirmation of acceptable human-label reliability. The exercise does not validate the automated outcome measures, and the first-round labels remain a provisional development reference pending a codebook-aligned repeat assessment and, ideally, independent second-coder evaluation."
).italic = True
shade_paragraph(p, LIGHT)

doc.add_heading("9. References", level=1)
doc.add_paragraph("Cohen, J. (1960). A coefficient of agreement for nominal scales. Educational and Psychological Measurement, 20(1), 37-46. https://doi.org/10.1177/001316446002000104")
doc.add_paragraph("Landis, J. R., & Koch, G. G. (1977). The measurement of observer agreement for categorical data. Biometrics, 33(1), 159-174. https://doi.org/10.2307/2529310")
doc.add_paragraph("McHugh, M. L. (2012). Interrater reliability: the kappa statistic. Biochemia Medica, 22(3), 276-282. https://doi.org/10.11613/BM.2012.031")

doc.add_heading("Appendix A. Reproducible statistics retained from the audit", level=1)
doc.add_paragraph(
    "For traceability, the following directional counts reproduce the supplied analysis. They are presented without assigning either round as ground truth."
)
rows2 = [
    ("coping_step", "25", "12", "0", "2"),
    ("professional_help", "34", "2", "1", "2"),
    ("social_support", "28", "6", "0", "5"),
    ("crisis_action", "10", "12", "0", "17"),
    ("follow_up", "0", "19", "0", "20"),
    ("surface_localisation", "35", "4", "0", "0"),
    ("verified_localisation", "18", "9", "0", "12"),
]
add_table(["Measure", "R1=1, R2=1", "R1=0, R2=1", "R1=1, R2=0", "R1=0, R2=0"], rows2, [2.0, 1.05, 1.05, 1.05, 1.05], 8)
doc.add_paragraph(
    "Actionability binary diagnostic: concordant positive = 30, R1 negative/R2 positive = 5, R1 positive/R2 negative = 1, concordant negative = 3; observed agreement = 84.6%; kappa = 0.421. This cross-scale diagnostic is not a substitute for weighted kappa on an invariant ordinal scale."
)

# Keep table rows together where feasible and avoid widows.
for p in doc.paragraphs:
    p.paragraph_format.widow_control = True

doc.core_properties.title = "Intra-Rater Repeat Coding Audit"
doc.core_properties.subject = "Validation supplement for MSc dissertation"
doc.core_properties.author = "Nor Lakrimdi"
doc.core_properties.comments = "Revised for methodological and statistical accuracy."
doc.save(OUT)
print(OUT)
